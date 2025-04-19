import streamlit as st
import openai
import os
import PyPDF2
from docx import Document

openai.api_key = "YOUR_API_KEY"

def summarize_text(text):
    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=f"Summarize the following text: {text}",
        max_tokens=150,
        n=1,
        stop=None,
        temperature=0.7,
    )
    return response.choices[0].text.strip()

def read_text_file(filepath):
    with open(filepath, 'r') as f:
        return f.read()

def read_pdf_file(filepath):
    with open(filepath, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text

def read_word_file(filepath):
    doc = Document(filepath)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

st.title("Generative AI Text Summarization Tool")

uploaded_file = st.file_uploader("Upload a document", type=["txt", "pdf", "docx"])

if uploaded_file is not None:
    file_type = uploaded_file.type
    file_name = uploaded_file.name

    if file_type == "text/plain":
        text = read_text_file(uploaded_file)
    elif file_type == "application/pdf":
        text = read_pdf_file(uploaded_file)
    elif file_name.endswith(".docx"):
        text = read_word_file(uploaded_file)
    else:
        st.error("Unsupported file type. Please upload a .txt, .pdf, or .docx file.")
        text = None

    if text:
        st.write("Original Text:")
        st.write(text)

        summary = summarize_text(text)

        st.write("Summary:")
        st.write(summary)

        # Download options
        st.download_button(
            label="Download Summary as TXT",
            data=summary,
            file_name="summary.txt",
            mime="text/plain",
        )

        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from io import BytesIO

        def create_pdf(text):
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            c.drawString(100, 750, text)
            c.save()
            buffer.seek(0)
            return buffer

        pdf_buffer = create_pdf(summary)

        st.download_button(
            label="Download Summary as PDF",
            data=pdf_buffer,
            file_name="summary.pdf",
            mime="application/pdf",
        )

        from docx import Document
        from io import BytesIO

        def create_word_doc(text):
            doc = Document()
            doc.add_paragraph(text)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        word_buffer = create_word_doc(summary)

        st.download_button(
            label="Download Summary as DOCX",
            data=word_buffer,
            file_name="summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
